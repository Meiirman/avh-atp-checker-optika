import traceback
from docx import Document

def read_cell(cell):
    """
    Функция для чтения содержимого ячейки, включая вложенные таблицы.
    """
    if cell.tables:
        # Если в ячейке есть вложенные таблицы, рекурсивно читаем их
        nested_tables = []
        for nested_table in cell.tables:
            nested_table_data = read_table(nested_table)
            nested_tables.append(nested_table_data)
        return nested_tables
    else:
        # В противном случае возвращаем текст из ячейки
        return cell.text


def read_table(table):
    """
    Функция для чтения содержимого таблицы, включая вложенные таблицы.
    """
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_data = read_cell(cell)
            row_data.append(cell_data)
        table_data.append(row_data)
    return table_data

def read_docx_tables(docx_file):
    doc = Document(docx_file)
    tables = []

    for table in doc.tables:
        table_data = read_table(table)
        tables.append(table_data)

    # doc = Document(docx_file)
    # result_table = []
    # for table in doc.tables:
    #     for row in table.rows:
    #         row_text = [cell.text.strip() for cell in row.cells]
    #         result_table.append(row_text)

    return tables





def parse_docx(docx_file):
    data = read_docx_tables(docx_file)
    print(data)
    print("____"*5)
    print(data)
    print("____"*5)


    # print(data[0][0][0][1])
    dd = data[1:]
    result = []
    for ddr in dd:
        if len(ddr) < 4:
            continue
        try:
            result.append(( float(ddr[0]), ddr[1], float(ddr[-3]), float(ddr[-2].replace(",",".").replace(" ","")) ))
        except:
            traceback.print_exc()
    
    data = []
    for i in result:
        data.append(( i[1], (i[0], i[2], i[3]) ))

    

    return {"message": "Успешно получил данные из DOCX", "data": data}


